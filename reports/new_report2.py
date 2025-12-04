from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_pet_hotel_report():
    doc = Document()

    # تنسيق عام
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ========== الصفحة الأولى ==========
    # العنوان الرئيسي
    header1 = doc.add_paragraph("МИНОБРНАУКИ РОССИИ")
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header2 = doc.add_paragraph("САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ")
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header3 = doc.add_paragraph("ЭЛЕКТРОТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ")
    header3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header4 = doc.add_paragraph("«ЛЭТИ» ИМ. В.И. УЛЬЯНОВА (ЛЕНИНА)")
    header4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header5 = doc.add_paragraph("Кафедра ИС")
    header5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # عنوان التقرير
    title = doc.add_paragraph("ОТЧЕТ")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.font.bold = True

    subtitle1 = doc.add_paragraph("по лабораторной работе №1 и 2")
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle2 = doc.add_paragraph("по дисциплине «Системы поддержки принятия решений»")
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    theme = doc.add_paragraph(
        "ТЕМА №52: БИЗНЕС-СЦЕНАРИЙ: ОРГАНИЗАЦИЯ РАБОТЫ ОТЕЛЯ ДЛЯ ДОМАШНИХ ЖИВОТНЫХ."
    )
    theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = theme.runs[0]
    run.font.bold = True

    subtitle3 = doc.add_paragraph("BPMN2.0/DMN/CAMUNDA8/ZEEBE")
    subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # معلومات الطالب
    student = doc.add_paragraph()
    student.add_run("Студент гр. 1378\t\t").font.bold = False
    student.add_run("Ханна М.Н.").font.bold = False

    teacher = doc.add_paragraph()
    teacher.add_run("Преподаватель\t\t").font.bold = False
    teacher.add_run("Васильев Н.В.").font.bold = False

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph("Санкт-Петербург")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2 = doc.add_paragraph("2025")
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # صفحة جديدة
    doc.add_page_break()

    # ========== 1. ЦЕЛЬ РАБОТЫ ==========
    heading1 = doc.add_heading("1. ЦЕЛЬ РАБОТЫ", level=1)

    p1 = doc.add_paragraph(
        "Целью данной работы является формализация бизнес-сценария «Отель для домашних животных» "
        "с использованием стандарта BPMN 2.0, с соблюдением использования только XOR и AND-шлюзов, "
        "подпроцессов и событий-таймеров, с последующим запуском процесса на движке Camunda Zeebe (л.р. 1)."
    )

    p2 = doc.add_paragraph(
        "Вторая цель – использование инструментария имитационного моделирования (BIMP) для анализа "
        "производительности процесса. Это включает определение среднего времени цикла, стоимости ресурсов "
        "и выявление характеристик выполнения процесса (л.р. 2)."
    )

    p3 = doc.add_paragraph(
        "Наконец, на основе полученных результатов симуляции, применяются методы анализа для выработки "
        "предложений по улучшению процесса и распределения ресурсов."
    )

    # ========== 2. СЦЕНАРИЙ ==========
    heading2 = doc.add_heading("2. СЦЕНАРИЙ «ОТЕЛЬ ДЛЯ ДОМАШНИХ ЖИВОТНЫХ»", level=1)

    doc.add_paragraph(
        "Процесс начинается с получения запроса на бронирование. Специалист по приему регистрирует бронь, "
        "после чего запускается проверка вакцинации (через DMN). Если вакцинация недействительна, "
        "бронь отклоняется."
    )

    doc.add_paragraph(
        "Когда питомец прибывает, начинается параллельный процесс регистрации:"
    )

    ul1 = doc.add_paragraph(
        "Ветеринарный осмотр: Ветеринар осматривает животное.", style="List Bullet"
    )
    ul2 = doc.add_paragraph(
        "Размещение в номере: Специалист по уходу готовит номер.", style="List Bullet"
    )

    doc.add_paragraph(
        "После завершения обеих задач запускается подпроцесс «Ежедневный уход», повторяющийся каждые 24 часа. "
        "Граничное событие-таймер «За день до выезда» инициирует отправку напоминания владельцу."
    )

    doc.add_paragraph(
        "По окончании срока пребывания, запускается процесс выписки: подготовка питомца, итоговый отчет, "
        "оплата и сбор отзыва. В конце процесс разделяется на два параллельных потока: завершение процесса "
        "для клиента и задача «Уборка номера» для хозяйственного отдела."
    )

    # ========== 3. ХОД РАБОТЫ ==========
    heading3 = doc.add_heading("3. ХОД РАБОТЫ (BPMN/DMN)", level=1)

    doc.add_heading("А. Диаграмма BPMN", level=2)

    doc.add_paragraph(
        "Диаграмма (файл `\\bpmn\\PetHotelProcess_v3.bpmn`) была спроектирована в Camunda Modeler "
        'для платформы Camunda 8 (Zeebe). Процесс разделен на 4 "дорожки" (Lanes) для демонстрации ролей:'
    )

    doc.add_paragraph("Специалист по приему", style="List Number")
    doc.add_paragraph("Ветеринар", style="List Number")
    doc.add_paragraph("Специалист по уходу", style="List Number")
    doc.add_paragraph("Хозяйственный отдел", style="List Number")

    doc.add_heading("Б. Таблицы решений (DMN)", level=2)

    doc.add_paragraph("Были использованы две DMN-таблицы для автоматизации решений:")

    doc.add_paragraph(
        "`vaccination_check`: Проверяет валидность вакцинации на основе дат.",
        style="List Number",
    )
    doc.add_paragraph(
        "`pet_compatibility_check`: Оценивает совместимость питомца для определения типа игр "
        "(групповые/индивидуальные).",
        style="List Number",
    )

    # ========== 4. НАСТРОЙКА СИМУЛЯЦИИ ==========
    heading4 = doc.add_heading("4. НАСТРОЙКА СИМУЛЯЦИИ (BIMP)", level=1)

    doc.add_paragraph(
        "Для анализа процесса в BIMP был подготовлен сценарий симуляции "
        "(файл `\\bpmn\\v3Bimp\\PetHotelProcess_v3_withInputs.bpmn`). "
        "Было определено 5 типов ресурсов и 2 расписания:"
    )

    doc.add_heading("Ресурсы (Resources):", level=3)

    # جدول الموارد
    table1 = doc.add_table(rows=6, cols=4)
    table1.style = "Light Grid Accent 1"

    # عناوين الجدول
    headers = ["Ресурс", "Кол-во", "Стоимость/час", "Расписание"]
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # بيانات الجدول
    resources_data = [
        ["Специалист по приему", "2", "20 USD", "Default"],
        ["Ветеринар", "1", "40 USD", "Default"],
        ["Специалист по уходу", "3", "18 USD", "24/7"],
        ["Хозяйственный отдел", "2", "16 USD", "Default"],
        ["Automated Service", "1", "0 USD", "24/7"],
    ]

    for i, row_data in enumerate(resources_data, start=1):
        for j, cell_data in enumerate(row_data):
            table1.rows[i].cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading("Расписания (Timetables):", level=3)

    doc.add_paragraph(
        "`Default`: Пн-Пт, 9:00 - 17:00 (для приема, ветеринара, хоз. отдела)."
    )
    doc.add_paragraph(
        "`24/7`: Все дни, круглосуточно (для специалистов по уходу и авто-сервисов)."
    )

    doc.add_heading("Шлюзы (Gateways):", level=3)

    doc.add_paragraph('"Вакцинация ОК?": 90% (Да), 10% (Нет).')
    doc.add_paragraph('"Совместим?": 70% (Да), 30% (Нет).')

    # ========== 5. АНАЛИЗ РЕЗУЛЬТАТОВ ==========
    doc.add_page_break()
    heading5 = doc.add_heading("5. АНАЛИЗ РЕЗУЛЬТАТОВ СИМУЛЯЦИИ И ОПТИМИЗАЦИЯ", level=1)

    doc.add_heading("А. Результаты симуляции (100 экземпляров процесса)", level=2)

    doc.add_paragraph(
        "На основе результатов симуляции (файлы `\\bpmn\\v3Bimp\\simulation_results.csv` "
        "и `\\bpmn\\v3Bimp\\PetHotelProcess_v3_withResults.bpmn`), были получены следующие показатели:"
    )

    doc.add_heading("1. Общие показатели процесса:", level=3)

    # جدول النتائج العامة
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = "Light Grid Accent 1"

    general_metrics = [
        ["Показатель", "Значение"],
        ["Количество экземпляров процесса", "100"],
        ["Среднее время цикла (Cycle Time)", "54 657 мин. ≈ 38 дней"],
        ["Среднее время работы (Duration)", "6 540 мин. ≈ 4.5 дня"],
        ["Средняя стоимость процесса", "67.45 USD"],
    ]

    for i, row_data in enumerate(general_metrics):
        for j, cell_data in enumerate(row_data):
            cell = table2.rows[i].cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading("2. Статистика по шлюзам принятия решений:", level=3)

    # جدول القرارات
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = "Light Grid Accent 1"

    decision_data = [
        ["Шлюз", "Путь", "Результат"],
        ["Вакцинация ОК?", "Да (принято)", "94 (94%)"],
        ["", "Нет (отклонено)", "6 (6%)"],
        ["Совместим? (из 94)", "Групповые игры", "66 (70.2%)"],
    ]

    for i, row_data in enumerate(decision_data):
        for j, cell_data in enumerate(row_data):
            cell = table3.rows[i].cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True

    # دمج الخلايا للصف الثاني
    table3.rows[2].cells[0].merge(table3.rows[1].cells[0])

    row_data_last = ["", "Индивидуальные игры", "28 (29.8%)"]
    for j, cell_data in enumerate(row_data_last):
        table3.add_row().cells[j].text = cell_data

    doc.add_paragraph()

    doc.add_heading("3. Загрузка ресурсов (Resource Utilization):", level=3)

    # جدول استخدام الموارد
    table4 = doc.add_table(rows=5, cols=2)
    table4.style = "Light Grid Accent 1"

    utilization_data = [
        ["Ресурс", "Загрузка"],
        ["Специалист по приему", "17.87%"],
        ["Ветеринар", "15.56%"],
        ["Специалист по уходу", "30.34%"],
        ["Хозяйственный отдел", "11.67%"],
    ]

    for i, row_data in enumerate(utilization_data):
        for j, cell_data in enumerate(row_data):
            cell = table4.rows[i].cells[j]
            cell.text = cell_data
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading("4. Время ожидания в задачах:", level=3)

    doc.add_paragraph(
        "Анализ показал, что среднее время ожидания (Waiting Time) для всех задач процесса "
        "составляет 0 минут. Это означает, что ресурсы были доступны немедленно при поступлении работы, "
        'и критических "узких мест" (bottlenecks) в текущей конфигурации не обнаружено.'
    )

    doc.add_heading("Б. Интерпретация результатов", level=2)

    doc.add_heading("1. Разница между Cycle Time и Duration:", level=3)

    doc.add_paragraph(
        "Ключевое наблюдение: существует значительная разница между средним временем цикла (38 дней) "
        "и средней продолжительностью работы (4.5 дня)."
    )

    doc.add_paragraph()
    p_explain = doc.add_paragraph()
    p_explain.add_run("Объяснение: ").font.bold = True
    p_explain.add_run(
        "Время цикла включает календарное время пребывания питомца в отеле (например, владелец "
        "оставил питомца на месяц). Продолжительность работы – это только время активных действий "
        "персонала (осмотр, уход, выписка и т.д.)."
    )

    doc.add_paragraph(
        'Анализ подпроцесса "Ежедневный уход" показал среднее время выполнения 27 281 минуту '
        "(≈ 19 дней), что коррелирует с длительностью пребывания питомцев."
    )

    doc.add_heading("2. Оценка загрузки ресурсов:", level=3)

    doc.add_paragraph(
        "Все ресурсы имеют низкую или умеренную загрузку (максимум 30.34% для специалиста по уходу). "
        "Это объясняется следующими факторами:"
    )

    doc.add_paragraph(
        "Низкая интенсивность поступления заявок: Симуляция использовала фиксированный интервал "
        "между заявками (7200 часов = 300 дней между экземплярами процесса).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Рабочее расписание: Ресурсы доступны только 8 часов в день (кроме специалистов по уходу), "
        "но фактическая работа занимает небольшую часть этого времени.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Отсутствие конкуренции за ресурсы: При данной интенсивности заявок несколько процессов "
        "редко выполняются параллельно.",
        style="List Bullet",
    )

    doc.add_heading('3. Отсутствие "узких мест":', level=3)

    doc.add_paragraph(
        "Нулевое время ожидания для всех задач указывает на то, что текущая конфигурация ресурсов "
        "полностью справляется с нагрузкой. Это положительный результат с точки зрения качества "
        "обслуживания, но также указывает на потенциал для оптимизации затрат."
    )

    doc.add_heading("В. Предложения по оптимизации", level=2)

    doc.add_paragraph(
        "Поскольку процесс работает без задержек, оптимизация должна быть направлена на "
        "повышение эффективности использования ресурсов и подготовку к росту нагрузки:"
    )

    doc.add_heading(
        "Оптимизация 1: Сокращение численности персонала (Cost Reduction)", level=3
    )

    p_opt1 = doc.add_paragraph()
    p_opt1.add_run("Применяемая эвристика: ").font.bold = True
    p_opt1.add_run('"Управление ресурсами" (Resource Management).')

    p_prob1 = doc.add_paragraph()
    p_prob1.add_run("Проблема: ").font.bold = True
    p_prob1.add_run(
        "Хозяйственный отдел загружен всего на 11.67%, но содержит 2 сотрудников. "
        "Специалисты по приему загружены на 17.87% (2 человека)."
    )

    p_sol1 = doc.add_paragraph()
    p_sol1.add_run("Решение: ").font.bold = True
    p_sol1.add_run(
        "Сократить хозяйственный отдел до 1 сотрудника. При необходимости (в пиковые периоды) "
        "использовать гибкий график или временный персонал. Для специалистов по приему можно "
        "рассмотреть совмещение функций с другими административными задачами отеля."
    )

    p_res1 = doc.add_paragraph()
    p_res1.add_run("Результат: ").font.bold = True
    p_res1.add_run(
        "Экономия на зарплатах без ухудшения качества обслуживания (т.к. время ожидания = 0)."
    )

    doc.add_heading(
        "Оптимизация 2: Подготовка к масштабированию (Scalability)", level=3
    )

    p_opt2 = doc.add_paragraph()
    p_opt2.add_run("Применяемая эвристика: ").font.bold = True
    p_opt2.add_run('"Планирование мощностей" (Capacity Planning).')

    p_prob2 = doc.add_paragraph()
    p_prob2.add_run("Проблема: ").font.bold = True
    p_prob2.add_run(
        "Текущая симуляция основана на очень низкой интенсивности заявок. "
        "При увеличении потока клиентов (например, в 10-20 раз) ресурсы могут стать недостаточными."
    )

    p_sol2 = doc.add_paragraph()
    p_sol2.add_run("Решение: ").font.bold = True
    p_sol2.add_run(
        "Провести повторную симуляцию с реалистичной интенсивностью (например, интервал между "
        'заявками 2-4 часа вместо 300 дней). Это позволит выявить реальные "узкие места" '
        "и определить оптимальное количество персонала."
    )

    doc.add_heading(
        "Оптимизация 3: Автоматизация административных задач (Automation)", level=3
    )

    p_opt3 = doc.add_paragraph()
    p_opt3.add_run("Применяемая эвристика: ").font.bold = True
    p_opt3.add_run('"Автоматизация задач" (Task Automation).')

    p_prob3 = doc.add_paragraph()
    p_prob3.add_run("Проблема: ").font.bold = True
    p_prob3.add_run(
        "Низкая загрузка специалистов по приему (17.87%) частично связана с рутинными "
        "задачами, которые можно автоматизировать."
    )

    p_sol3 = doc.add_paragraph()
    p_sol3.add_run("Решение: ").font.bold = True
    p_sol3.add_run(
        "Внедрить онлайн-систему бронирования с автоматической проверкой вакцинации (уже частично "
        "реализовано через DMN). Автоматизировать отправку отчетов владельцам (Task_OwnerUpdate) "
        "через email/мессенджеры. Использовать электронный документооборот для договоров."
    )

    p_res3 = doc.add_paragraph()
    p_res3.add_run("Результат: ").font.bold = True
    p_res3.add_run(
        "Высвобождение времени персонала для более важных задач (например, взаимодействие с клиентами). "
        "Снижение вероятности ошибок."
    )

    doc.add_heading(
        "Оптимизация 4: Оптимизация расписания (Schedule Optimization)", level=3
    )

    p_opt4 = doc.add_paragraph()
    p_opt4.add_run("Применяемая эвристика: ").font.bold = True
    p_opt4.add_run('"Адаптация расписания к спросу" (Demand-based Scheduling).')

    p_prob4 = doc.add_paragraph()
    p_prob4.add_run("Наблюдение: ").font.bold = True
    p_prob4.add_run(
        "Специалисты по уходу работают 24/7, что необходимо для обслуживания животных. "
        "Однако остальной персонал работает только 9-17, что может вызвать задержки при "
        "увеличении потока клиентов."
    )

    p_sol4 = doc.add_paragraph()
    p_sol4.add_run("Решение: ").font.bold = True
    p_sol4.add_run(
        "Анализировать паттерны прибытия/выбытия клиентов. Если большинство выписок происходит "
        "утром (7-9), можно добавить утреннюю смену для специалистов по приему. "
        "Для ветеринара можно ввести дежурства в выходные дни."
    )

    # ========== 6. ВЫВОДЫ ==========
    doc.add_page_break()
    heading6 = doc.add_heading("6. ВЫВОДЫ", level=1)

    doc.add_paragraph(
        "Бизнес-процесс «Отель для домашних животных» был успешно смоделирован как исполняемый "
        "процесс на Camunda Zeebe с использованием стандарта BPMN 2.0 и таблиц решений DMN (л.р. 1)."
    )

    doc.add_paragraph(
        "Имитационное моделирование с использованием BIMP показало себя как эффективный инструмент "
        "для анализа производительности процесса (л.р. 2). Основные результаты симуляции (100 экземпляров):"
    )

    doc.add_paragraph(
        "Среднее время цикла составило 38 дней, что отражает типичную продолжительность "
        "пребывания питомцев в отеле.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Средняя продолжительность активной работы персонала – 4.5 дня на один экземпляр процесса.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Средняя стоимость обслуживания одного питомца – 67.45 USD.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Уровень отказов по вакцинации составил 6%, что соответствует ожиданиям.",
        style="List Bullet",
    )
    doc.add_paragraph(
        'Отсутствие времени ожидания во всех задачах указывает на отсутствие "узких мест" '
        "при текущей интенсивности заявок.",
        style="List Bullet",
    )

    doc.add_paragraph(
        "Анализ загрузки ресурсов выявил значительный потенциал для оптимизации:"
    )

    doc.add_paragraph(
        "Низкая загрузка персонала (11-30%) при низкой интенсивности заявок открывает "
        "возможности для сокращения затрат.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Рекомендуется провести дополнительную симуляцию с реалистичной интенсивностью "
        "поступления заявок для определения оптимального штата.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Автоматизация рутинных административных задач может высвободить ресурсы персонала.",
        style="List Bullet",
    )

    doc.add_paragraph(
        "Применение методов анализа и эвристик оптимизации позволило разработать практические "
        "рекомендации по улучшению процесса:"
    )

    doc.add_paragraph(
        "Управление численностью персонала на основе фактической загрузки.",
        style="List Number",
    )
    doc.add_paragraph(
        "Подготовка к масштабированию через планирование мощностей.",
        style="List Number",
    )
    doc.add_paragraph(
        "Автоматизация административных задач для повышения эффективности.",
        style="List Number",
    )
    doc.add_paragraph(
        "Адаптация расписания персонала к паттернам спроса.", style="List Number"
    )

    doc.add_paragraph()

    doc.add_paragraph(
        "Данная работа продемонстрировала практическую ценность имитационного моделирования "
        "для принятия обоснованных решений по управлению бизнес-процессами. Полученные результаты "
        "могут быть использованы для реального внедрения системы управления отелем для домашних животных."
    )

    # حفظ المستند
    doc.save("Отчет_Отель_для_животных_ИСПРАВЛЕННЫЙ.docx")
    print("✅ Отчет успешно создан: Отчет_Отель_для_животных_ИСПРАВЛЕННЫЙ.docx")


# تشغيل الدالة
if __name__ == "__main__":
    create_pet_hotel_report()
